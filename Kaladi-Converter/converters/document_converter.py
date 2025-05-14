# from .base_converter import BaseConverter
# from docx2pdf import convert
# from docx import Document
# import io
# import PyPDF2
# import streamlit as st
# from pdf2image import convert_from_bytes
# import pythoncom
# from typing import Union
# import warnings
# import tempfile
# import os

# class DocumentConverter(BaseConverter):
#     def __init__(self):
#         super().__init__()
#         self.title = "Document Converter"
#         self.icon = "📝"
#         self.description = "Convert between various document formats including DOC, DOCX, PDF, TXT, and more."
#         self.supported_formats = {
#             'doc': ['pdf', 'txt'],
#             'docx': ['pdf', 'txt'],
#             'pdf': ['docx', 'txt', 'png', 'jpg'],
#             'txt': ['docx', 'pdf'],
#             'odt': ['docx', 'pdf'],
#             'rtf': ['docx', 'pdf']
#         }
#         warnings.filterwarnings("ignore", category=UserWarning)

#     def _validate_pdf(self, file_bytes: bytes) -> PyPDF2.PdfReader:
#         """Validate PDF file and return reader object."""
#         try:
#             # Handle both PyPDF2 v1 and v2 exceptions
#             try:
#                 # PyPDF2 v2
#                 pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
#                 if len(pdf_reader.pages) == 0:
#                     raise ValueError("PDF appears to be empty")
#                 return pdf_reader
#             except Exception as e:
#                 # PyPDF2 v1 fallback
#                 pdf_reader = PyPDF2.PdfFileReader(io.BytesIO(file_bytes))
#                 if pdf_reader.getNumPages() == 0:
#                     raise ValueError("PDF appears to be empty")
#                 return pdf_reader
#         except Exception as e:
#             raise ValueError(f"Invalid PDF file: {str(e)}")

#     def convert(self, input_file: Union[io.BytesIO, st.runtime.uploaded_file_manager.UploadedFile], 
#                 output_format: str) -> bytes:
#         input_format = input_file.name.split('.')[-1].lower()
#         file_bytes = input_file.getvalue()
        
#         # Check if conversion is needed (input == output)
#         if input_format == output_format.lower():
#             raise ValueError(f"Source file is already in {output_format.upper()} format. No conversion needed.")

#         try:
#             if input_format in ['doc', 'docx']:
#                 if output_format == 'pdf':
#                     # Create temporary files
#                     with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
#                         temp_docx.write(file_bytes)
#                         temp_docx_path = temp_docx.name
                    
#                     with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
#                         temp_pdf_path = temp_pdf.name
                    
#                     try:
#                         # Initialize COM for docx2pdf
#                         pythoncom.CoInitialize()
#                         convert(temp_docx_path, temp_pdf_path)
                        
#                         # Read the converted PDF
#                         with open(temp_pdf_path, 'rb') as f:
#                             pdf_bytes = f.read()
                        
#                         return pdf_bytes
#                     finally:
#                         pythoncom.CoUninitialize()
#                         # Clean up temporary files
#                         try:
#                             os.unlink(temp_docx_path)
#                             os.unlink(temp_pdf_path)
#                         except:
#                             pass
                
#                 elif output_format == 'txt':
#                     doc = Document(io.BytesIO(file_bytes))
#                     text = "\n".join([para.text for para in doc.paragraphs])
#                     return text.encode('utf-8')

#             elif input_format == 'pdf':
#                 if output_format in ['docx', 'txt']:
#                     pdf_reader = self._validate_pdf(file_bytes)
#                     text = ""
                    
#                     # Handle both PyPDF2 versions
#                     pages = pdf_reader.pages if hasattr(pdf_reader, 'pages') else [pdf_reader.getPage(i) for i in range(pdf_reader.getNumPages())]
                    
#                     for page in pages:
#                         page_text = page.extract_text() if hasattr(page, 'extract_text') else page.extractText()
#                         if page_text:
#                             text += page_text + "\n"
                    
#                     if output_format == 'txt':
#                         return text.encode('utf-8')
#                     else:
#                         doc = Document()
#                         doc.add_paragraph(text)
#                         output = io.BytesIO()
#                         doc.save(output)
#                         return output.getvalue()
                
#                 elif output_format in ['png', 'jpg']:
#                     images = convert_from_bytes(file_bytes, first_page=1, last_page=1)
#                     if not images:
#                         raise ValueError("No images found in PDF")
                    
#                     output = io.BytesIO()
#                     if output_format == 'jpg':
#                         images[0].convert('RGB').save(output, format='JPEG', quality=95)
#                     else:
#                         images[0].save(output, format='PNG')
#                     return output.getvalue()

#             elif input_format == 'txt':
#                 text = file_bytes.decode('utf-8')
                
#                 if output_format == 'docx':
#                     doc = Document()
#                     doc.add_paragraph(text)
#                     output = io.BytesIO()
#                     doc.save(output)
#                     return output.getvalue()
                
#                 elif output_format == 'pdf':
#                     # Create temporary files for txt to docx to pdf conversion
#                     with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
#                         doc = Document()
#                         doc.add_paragraph(text)
#                         doc.save(temp_docx)
#                         temp_docx_path = temp_docx.name
                    
#                     with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
#                         temp_pdf_path = temp_pdf.name
                    
#                     try:
#                         pythoncom.CoInitialize()
#                         convert(temp_docx_path, temp_pdf_path)
                        
#                         with open(temp_pdf_path, 'rb') as f:
#                             pdf_bytes = f.read()
                        
#                         return pdf_bytes
#                     finally:
#                         pythoncom.CoUninitialize()
#                         try:
#                             os.unlink(temp_docx_path)
#                             os.unlink(temp_pdf_path)
#                         except:
#                             pass

#             raise NotImplementedError(f"Conversion from {input_format} to {output_format} is not currently supported.")

#         except Exception as e:
#             raise ValueError(f"Conversion failed: {str(e)}")


from .base_converter import BaseConverter
from docx2pdf import convert
from docx import Document
import io
import PyPDF2
import streamlit as st
from pdf2image import convert_from_bytes
import pythoncom
import comtypes.client
from typing import Union
import warnings
import tempfile
import os
import subprocess
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import logging

# Setup logging for debugging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

class DocumentConverter(BaseConverter):
    def __init__(self):
        super().__init__()
        self.title = "Document Converter"
        self.icon = "📝"
        self.description = "Convert between various document formats including DOC, DOCX, PDF, TXT, and more."
        self.supported_formats = {
            'doc': ['pdf', 'txt'],
            'docx': ['pdf', 'txt'],
            'pdf': ['docx', 'txt', 'png', 'jpg'],
            'txt': ['docx', 'pdf'],
            'odt': ['docx', 'pdf'],
            'rtf': ['docx', 'pdf']
        }
        warnings.filterwarnings("ignore", category=UserWarning)

    def _validate_pdf(self, file_bytes: bytes) -> PyPDF2.PdfReader:
        """Validate PDF file and return reader object."""
        try:
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                if len(pdf_reader.pages) == 0:
                    raise ValueError("PDF appears to be empty")
                return pdf_reader
            except Exception:
                pdf_reader = PyPDF2.PdfFileReader(io.BytesIO(file_bytes))
                if pdf_reader.getNumPages() == 0:
                    raise ValueError("PDF appears to be empty")
                return pdf_reader
        except Exception as e:
            raise ValueError(f"Invalid PDF file: {str(e)}")

    def _convert_docx_to_pdf_docx2pdf(self, input_path: str, output_path: str) -> bytes:
        """Convert DOCX to PDF using docx2pdf."""
        try:
            pythoncom.CoInitialize()
            convert(input_path, output_path)
            with open(output_path, 'rb') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"docx2pdf conversion failed: {str(e)}")
        finally:
            pythoncom.CoUninitialize()

    def _convert_docx_to_pdf_comtypes(self, input_path: str, output_path: str) -> bytes:
        """Convert DOCX to PDF using comtypes."""
        try:
            word = comtypes.client.CreateObject("Word.Application")
            doc = word.Documents.Open(os.path.abspath(input_path))
            doc.SaveAs(os.path.abspath(output_path), FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
            with open(output_path, 'rb') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"comtypes conversion failed: {str(e)}")

    def _convert_docx_to_pdf_libreoffice(self, input_path: str, output_path: str) -> bytes:
        """Convert DOCX to PDF using LibreOffice."""
        try:
            subprocess.run([
                "soffice", "--headless", "--convert-to", "pdf",
                input_path, "--outdir", os.path.dirname(output_path)
            ], check=True, capture_output=True, text=True)
            output_pdf = os.path.join(os.path.dirname(output_path), os.path.basename(input_path).rsplit('.', 1)[0] + '.pdf')
            with open(output_pdf, 'rb') as f:
                pdf_bytes = f.read()
            os.unlink(output_pdf)  # Clean up
            return pdf_bytes
        except subprocess.CalledProcessError as e:
            raise ValueError(f"LibreOffice conversion failed: {e.stderr}")
        except FileNotFoundError:
            raise ValueError("LibreOffice not installed or 'soffice' not in PATH")

    def _convert_docx_to_pdf_reportlab(self, input_path: str, output_path: str) -> bytes:
        """Convert DOCX to PDF using python-docx and reportlab (basic formatting)."""
        try:
            if not os.path.exists(input_path):
                raise FileNotFoundError(f"Input file not found at {input_path}")
            doc = Document(input_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            pdf = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = [Paragraph(text, styles["Normal"])]
            pdf.build(story)
            with open(output_path, 'rb') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"reportlab conversion failed: {str(e)}")

    def convert(self, input_file: Union[io.BytesIO, st.runtime.uploaded_file_manager.UploadedFile], 
                output_format: str) -> bytes:
        input_format = input_file.name.split('.')[-1].lower()
        file_bytes = input_file.getvalue()
        
        if input_format == output_format.lower():
            raise ValueError(f"Source file is already in {output_format.upper()} format. No conversion needed.")

        # Create temporary files outside try block to persist across fallbacks
        temp_docx = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
        temp_pdf = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
        temp_docx_path = temp_docx.name
        temp_pdf_path = temp_pdf.name

        try:
            # Write input file to temporary DOCX
            temp_docx.write(file_bytes)
            temp_docx.close()  # Close to ensure file is written

            if input_format in ['doc', 'docx']:
                if output_format == 'pdf':
                    logging.debug(f"Attempting DOCX to PDF conversion for {input_file.name}")
                    # Try conversion methods in order
                    for method, func in [
                        ("docx2pdf", self._convert_docx_to_pdf_docx2pdf),
                        ("comtypes", self._convert_docx_to_pdf_comtypes),
                        ("LibreOffice", self._convert_docx_to_pdf_libreoffice),
                        ("reportlab", self._convert_docx_to_pdf_reportlab)
                    ]:
                        try:
                            logging.debug(f"Trying {method} conversion")
                            return func(temp_docx_path, temp_pdf_path)
                        except Exception as e:
                            logging.error(f"{method} failed: {str(e)}")
                            continue
                    raise ValueError("All conversion methods failed")
                
                elif output_format == 'txt':
                    doc = Document(io.BytesIO(file_bytes))
                    text = "\n".join([para.text for para in doc.paragraphs])
                    return text.encode('utf-8')

            elif input_format == 'pdf':
                if output_format in ['docx', 'txt']:
                    pdf_reader = self._validate_pdf(file_bytes)
                    text = ""
                    pages = pdf_reader.pages if hasattr(pdf_reader, 'pages') else [pdf_reader.getPage(i) for i in range(pdf_reader.getNumPages())]
                    for page in pages:
                        page_text = page.extract_text() if hasattr(page, 'extract_text') else page.extractText()
                        if page_text:
                            text += page_text + "\n"
                    if output_format == 'txt':
                        return text.encode('utf-8')
                    else:
                        doc = Document()
                        doc.add_paragraph(text)
                        output = io.BytesIO()
                        doc.save(output)
                        return output.getvalue()
                
                elif output_format in ['png', 'jpg']:
                    images = convert_from_bytes(file_bytes, first_page=1, last_page=1)
                    if not images:
                        raise ValueError("No images found in PDF")
                    output = io.BytesIO()
                    if output_format == 'jpg':
                        images[0].convert('RGB').save(output, format='JPEG', quality=95)
                    else:
                        images[0].save(output, format='PNG')
                    return output.getvalue()

            elif input_format == 'txt':
                text = file_bytes.decode('utf-8')
                if output_format == 'docx':
                    doc = Document()
                    doc.add_paragraph(text)
                    output = io.BytesIO()
                    doc.save(output)
                    return output.getvalue()
                elif output_format == 'pdf':
                    temp_docx.write(file_bytes)
                    temp_docx.close()
                    for method, func in [
                        ("docx2pdf", self._convert_docx_to_pdf_docx2pdf),
                        ("comtypes", self._convert_docx_to_pdf_comtypes),
                        ("LibreOffice", self._convert_docx_to_pdf_libreoffice),
                        ("reportlab", self._convert_docx_to_pdf_reportlab)
                    ]:
                        try:
                            logging.debug(f"Trying {method} conversion for TXT to PDF")
                            return func(temp_docx_path, temp_pdf_path)
                        except Exception as e:
                            logging.error(f"{method} failed: {str(e)}")
                            continue
                    raise ValueError("All conversion methods failed")

            raise NotImplementedError(f"Conversion from {input_format} to {output_format} is not currently supported.")

        except Exception as e:
            raise ValueError(f"Conversion failed: {str(e)}")
        finally:
            # Clean up temporary files
            try:
                if os.path.exists(temp_docx_path):
                    os.unlink(temp_docx_path)
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)
            except Exception as e:
                logging.warning(f"Failed to clean up temporary files: {str(e)}")